# CooDX11 - The Pigeon Word Counterpart GUI
# Developed by WingTech (c) 2011
# Version 2.0.2

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import random
import time
# Import the python-docx library for .docx functionality
try:
    from docx import Document
    from docx.shared import Inches # Not strictly used for text, but often part of docx imports
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx library not found. .docx saving/opening will be disabled.")
    print("Install it with: pip install python-docx")


# --- CooDX11 Core Logic (from previous version, adapted for GUI) ---

def pigeon_sound():
    """Returns a random pigeon sound."""
    sounds = ["Coo!", "Coo-coo!", "Rrrruh...", "Whirr!"]
    return random.choice(sounds)

def simplify_word(word):
    """
    Simplifies a word to a pigeon-like equivalent.
    This is a very basic simplification, mimicking limited vocabulary.
    """
    word = word.lower()
    if "food" in word or "eat" in word or "seed" in word or "yum" in word:
        return "Grain!"
    if "fly" in word or "wing" in word or "air" or "sky" in word:
        return "Sky!"
    if "home" in word or "nest" in word or "roost" or "place" in word:
        return "Perch!"
    if "friend" in word or "pal" or "mate" or "fellow" in word:
        return "Flock!"
    if "danger" in word or "cat" or "hawk" or "threat" or "scary" in word:
        return "Threat!"
    if "water" in word or "drink" or "liquid" in word:
        return "Drip!"
    if "hello" in word or "hi" or "greet" or "hey" in word:
        return "Coo?" # A questioning coo for greeting
    if "goodbye" in word or "bye" or "farewell" in word:
        return "Flap!" # Sound of departure
    if "yes" in word or "ok" in word or "affirmative" in word:
        return "Nod!"
    if "no" in word or "negative" in word:
        return "Shake!"
    if "love" in word or "care" in word or "like" in word:
        return "Preen!"
    if "walk" in word or "move" in word or "go" in word:
        return "Waddle!"
    return pigeon_sound() # Default to a generic coo

def translate_to_pigeon(human_text):
    """
    Translates human text into a pigeon-like sequence.
    It breaks down the text, simplifies words, and adds repetitive sounds.
    """
    words = human_text.split()
    pigeon_output = []

    if not words:
        return f"{pigeon_sound()} {pigeon_sound()} {pigeon_sound()}" # Empty input gets just coos

    for i, word in enumerate(words):
        p_word = simplify_word(word)
        pigeon_output.append(p_word)

        if p_word != pigeon_sound():
            if random.random() < 0.7:
                pigeon_output.append(pigeon_sound())
        
        if i % 3 == 0 and i != 0:
            if random.random() < 0.5:
                pigeon_output.append("Flap-flurry!")
            else:
                pigeon_output.append("...")

    if len(pigeon_output) < 5 and len(words) > 0: # Ensure some length for short inputs
        pigeon_output.extend([pigeon_sound(), pigeon_sound()])

    return " ".join(pigeon_output).replace("  ", " ").strip()

def reverse_translate_pigeon(pigeon_text):
    """
    Attempts to reverse translate pigeon-like text to human text.
    This is highly interpretive due to the simplification process.
    """
    pigeon_text = pigeon_text.lower()
    human_output_parts = []
    
    # Simple mapping of pigeon words back to human concepts
    mapping = {
        "grain!": "food", "sky!": "fly", "perch!": "home", "flock!": "friend",
        "threat!": "danger", "drip!": "water", "coo?": "hello", "flap!": "goodbye",
        "nod!": "yes", "shake!": "no", "preen!": "love", "waddle!": "walk"
    }

    words = pigeon_text.replace("coo-coo!", "coo!").replace("whirr!", "coo!").split()
    
    for word in words:
        # Remove punctuation if present
        clean_word = word.strip(".,!?;:").lower()
        if clean_word in mapping:
            human_output_parts.append(mapping[clean_word])
        elif "coo" in clean_word:
            human_output_parts.append("general observation")
        elif "flurry" in clean_word:
            human_output_parts.append("activity")
        elif "..." in clean_word:
            human_output_parts.append("pause")
        elif "rrrruh" in clean_word:
            human_output_parts.append("hesitation")
        else:
            human_output_parts.append("[unclear]") # For words not directly mapped

    return " ".join(human_output_parts).capitalize() + "."


# --- GUI Application ---

class CooDX11App:
    def __init__(self, master):
        self.master = master
        master.title("CooDX11 - WingTech 2011")
        master.geometry("800x600")
        master.resizable(True, True)
        
        # Apply a classic 2011-style theme
        style = ttk.Style()
        style.theme_use('vista') # Or 'xpnative', 'clam', 'alt' for a more classic feel
        style.configure("TButton", font=("Arial", 10), padding=5, relief="raised")
        style.configure("TLabel", font=("Arial", 10))
        style.configure("TFrame", background="#E0E0E0", relief="groove") # Light gray background
        
        # Define a style for the paper frame to allow background color
        style.configure("Paper.TFrame", background="#FFFFFF")

        # Menu Bar
        self.menubar = tk.Menu(master, bg="#F0F0F0", fg="#333333", activebackground="#D0D0D0", activeforeground="black")
        master.config(menu=self.menubar)

        # File Menu
        file_menu = tk.Menu(self.menubar, tearoff=0, bg="#F0F0F0", fg="#333333")
        self.menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New", command=self.new_document)
        file_menu.add_command(label="Open...", command=self.open_document)
        file_menu.add_command(label="Save", command=self.save_document)
        file_menu.add_command(label="Save As...", command=self.save_document_as)
        file_menu.add_separator()
        file_menu.add_command(label="Print...", command=self.print_document)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=master.quit)

        # Edit Menu (Placeholder)
        edit_menu = tk.Menu(self.menubar, tearoff=0, bg="#F0F0F0", fg="#333333")
        self.menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Cut", command=lambda: self.text_area.event_generate("<<Cut>>"))
        edit_menu.add_command(label="Copy", command=lambda: self.text_area.event_generate("<<Copy>>"))
        edit_menu.add_command(label="Paste", command=lambda: self.text_area.event_generate("<<Paste>>"))

        # Options Menu (Placeholder)
        options_menu = tk.Menu(self.menubar, tearoff=0, bg="#F0F0F0", fg="#333333")
        self.menubar.add_cascade(label="Options", menu=options_menu)
        options_menu.add_command(label="Settings...", command=self.show_settings)

        # Help Menu (Placeholder)
        help_menu = tk.Menu(self.menubar, tearoff=0, bg="#F0F0F0", fg="#333333")
        self.menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About CooDX11", command=self.show_about)

        # Main content frame (simulates the application body)
        main_frame = ttk.Frame(master, padding="10", relief="sunken")
        main_frame.pack(expand=True, fill="both")

        # Control Frame (buttons for translation/modes)
        control_frame = ttk.Frame(main_frame, padding="10", relief="raised")
        control_frame.pack(side="top", fill="x", pady=(0, 10))

        self.human_to_pigeon_button = ttk.Button(control_frame, text="Human to Pigeon ➡️", command=self.human_to_pigeon)
        self.human_to_pigeon_button.pack(side="left", padx=5, pady=5)

        self.pigeon_to_human_button = ttk.Button(control_frame, text="Pigeon to Human ⬅️", command=self.pigeon_to_human)
        self.pigeon_to_human_button.pack(side="left", padx=5, pady=5)

        self.just_pigeon_button = ttk.Button(control_frame, text="Generate Pigeon Talk 🐦", command=self.generate_just_pigeon)
        self.just_pigeon_button.pack(side="left", padx=5, pady=5)
        
        # Paper Area Frame (simulates A4 paper)
        # Apply the custom style here
        paper_frame = ttk.Frame(main_frame, padding="20", relief="sunken", borderwidth=2,
                                style="Paper.TFrame") 
        paper_frame.pack(expand=True, fill="both", padx=10, pady=10)
        
        # Using a scrolled text widget for the "paper" to allow scrolling
        self.text_area = scrolledtext.ScrolledText(paper_frame, wrap="word", width=70, height=25,
                                                   font=("Courier New", 12),
                                                   bg="#F8F8F8", fg="#333333",
                                                   insertbackground="black",
                                                   relief="flat", borderwidth=0)
        self.text_area.pack(expand=True, fill="both")
        
        # Status Bar (typical of 2011 apps)
        self.status_bar = ttk.Label(master, text="Ready.", relief="sunken", anchor="w")
        self.status_bar.pack(side="bottom", fill="x", ipady=2)

    def new_document(self):
        """Clears the text area for a new document."""
        if messagebox.askyesno("New Document", "Do you want to clear the current text?"):
            self.text_area.delete(1.0, tk.END)
            self.status_bar.config(text="New document created.")

    def open_document(self):
        """Simulates opening a document, now with .docx support."""
        file_path = filedialog.askopenfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"),
                       ("Word Documents", "*.docx"),
                       ("All files", "*.*")]
        )
        if file_path:
            try:
                if file_path.lower().endswith('.docx'):
                    if DOCX_AVAILABLE:
                        doc = Document(file_path)
                        full_text = []
                        for para in doc.paragraphs:
                            full_text.append(para.text)
                        content = "\n".join(full_text)
                        self.text_area.delete(1.0, tk.END)
                        self.text_area.insert(1.0, content)
                        self.status_bar.config(text=f"Opened: {file_path}")
                    else:
                        messagebox.showerror("Error", "python-docx library not available for opening .docx files. Please install it.")
                        self.status_bar.config(text="Failed to open document (docx support missing).")
                else: # Assume text file for other extensions or default
                    with open(file_path, "r", encoding="utf-8") as file: # Added encoding for broader compatibility
                        content = file.read()
                        self.text_area.delete(1.0, tk.END)
                        self.text_area.insert(1.0, content)
                    self.status_bar.config(text=f"Opened: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Could not open file: {e}")
                self.status_bar.config(text="Failed to open document.")

    def save_document(self):
        """Simulates saving a document."""
        # In a real app, you'd track the current file path. For simplicity,
        # save acts like save_as here.
        self.save_document_as() 

    def save_document_as(self):
        """Simulates saving a document to a new location, now including .docx."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"),
                       ("Word Documents", "*.docx"),
                       ("All files", "*.*")]
        )
        if file_path:
            try:
                if file_path.lower().endswith('.docx'):
                    if DOCX_AVAILABLE:
                        document = Document()
                        document.add_paragraph(self.text_area.get(1.0, tk.END))
                        document.save(file_path)
                        self.status_bar.config(text=f"Saved to: {file_path} (DOCX)")
                    else:
                        messagebox.showerror("Error", "python-docx library not available for saving .docx files. Please install it.")
                        self.status_bar.config(text="Failed to save document (docx support missing).")
                else: # Default to text file
                    with open(file_path, "w", encoding="utf-8") as file: # Added encoding for broader compatibility
                        file.write(self.text_area.get(1.0, tk.END))
                    self.status_bar.config(text=f"Saved to: {file_path} (TXT)")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save file: {e}")
                self.status_bar.config(text="Failed to save document.")

    def print_document(self):
        """Simulates a print action."""
        # In a real application, this would interact with a printer driver.
        # For this simulation, we'll just show a message.
        messagebox.showinfo("Print Document", "Sending document to virtual printer... (Print job simulated!)")
        self.status_bar.config(text="Print job simulated.")

    def human_to_pigeon(self):
        """Translates text from human to pigeon."""
        human_text = self.text_area.get(1.0, tk.END).strip()
        if not human_text:
            messagebox.showwarning("Input Empty", "Please enter some human text to translate.")
            return

        self.status_bar.config(text="Translating to pigeon...")
        # Simulate processing time
        self.master.update_idletasks()
        time.sleep(0.1) # Brief delay

        pigeon_message = translate_to_pigeon(human_text)
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, pigeon_message)
        self.status_bar.config(text="Translation complete: Human to Pigeon.")

    def pigeon_to_human(self):
        """Translates text from pigeon to human."""
        pigeon_text = self.text_area.get(1.0, tk.END).strip()
        if not pigeon_text:
            messagebox.showwarning("Input Empty", "Please enter some pigeon text to translate.")
            return

        self.status_bar.config(text="Translating to human (best effort)...")
        self.master.update_idletasks()
        time.sleep(0.1)

        human_message = reverse_translate_pigeon(pigeon_text)
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, human_message)
        self.status_bar.config(text="Translation complete: Pigeon to Human (interpretive).")

    def generate_just_pigeon(self):
        """Generates random pigeon-like dialogue."""
        pigeon_phrases = [
            "Coo! Grain! Coo-coo! Flap-flurry!",
            "Perch! Coo? Threat! ... Coo!",
            "Flock! Coo! Sky! Whirr!",
            "Drip! Coo-coo! Grain! Coo?"
        ]
        
        # Generate a longer random pigeon text
        generated_text = ""
        for _ in range(random.randint(3, 7)): # A few paragraphs of pigeon talk
            generated_text += random.choice(pigeon_phrases) + " " * random.randint(1, 3)
            generated_text += pigeon_sound() + " " * random.randint(1, 2)
            if random.random() < 0.3:
                generated_text += "\n\n" # Simulate new paragraphs
        
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, generated_text.strip())
        self.status_bar.config(text="Generated random pigeon talk.")

    def show_settings(self):
        """Placeholder for settings dialog."""
        messagebox.showinfo("Settings", "CooDX11 Settings (Feature coming soon!)")

    def show_about(self):
        """Shows the About dialog."""
        about_text = ("CooDX11 - The Pigeon Word Counterpart\n"
                      "Version 2.0.2 (GUI Edition)\n" # Updated version number
                      "Developed by WingTech (c) 2011\n\n"
                      "Translating human thoughts to pigeon wisdom, and back!")
        messagebox.showinfo("About CooDX11", about_text)


if __name__ == "__main__":
    root = tk.Tk()
    app = CooDX11App(root)
    root.mainloop()

